#!/usr/bin/env python3
"""
build_doc.py — Assemble the Blazepod deep-dive .docx deliverable.

WHAT IT DOES
    Reads the structured research JSON gathered by the agent during the workflow,
    and renders a well-structured Word document explaining:
      1. What Blazepod is (overview)
      2. The hardware, explained for a reader with no hardware knowledge
      3. How the software works (app, drills, BLE protocol, metrics)
      4. Build notes mapping each subsystem to a clone-able part

USAGE
    python tools/build_doc.py --research .tmp/blazepod_research.json --out export/blazepod-deep-dive.docx

INPUT SCHEMA (.tmp/blazepod_research.json)
    {
      "title": "...",
      "generated_at": "2026-...",
      "sources": [{"label": "...", "url": "..."}],
      "overview": "long markdown-ish string",
      "hardware": [
        {"subsystem": "Battery", "blazepod_spec": "...", "plain_explanation": "...",
         "clone_part": "3.7V 600mAh Li-Po", "confidence": "confirmed|inferred|unknown"}
      ],
      "software": [
        {"topic": "How pods pair", "explanation": "...", "confidence": "..."}
      ],
      "sources": [...]
    }
"""

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.stderr.write(
        "Missing dependency. Install with:\n  pip install -r requirements.txt\n"
    )
    raise

CONFIDENCE_COLORS = {
    "confirmed": RGBColor(0x1B, 0x7A, 0x43),  # green
    "inferred": RGBColor(0xB8, 0x6A, 0x00),   # amber
    "unknown": RGBColor(0xA1, 0x1D, 0x1D),    # red
}


def _load(research_path):
    with Path(research_path).open(encoding="utf-8") as f:
        return json.load(f)


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def _add_paragraph(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p


def _add_confidence(doc, confidence):
    """Render a small confidence tag after a subsystem heading."""
    if not confidence:
        return
    p = doc.add_paragraph()
    run = p.add_run(f"Confidence: {confidence}")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = CONFIDENCE_COLORS.get(
        confidence, RGBColor(0x66, 0x66, 0x66)
    )


def _render_markdownish(doc, text):
    """Render a loose markdown-like block: split on blank lines into paragraphs,
    and treat lines starting with '- ' or '* ' as bullets."""
    if not text:
        _add_paragraph(doc, "(no data)")
        return
    blocks = re_split_blank(text)
    for block in blocks:
        lines = [ln for ln in block.splitlines() if ln.strip()]
        if not lines:
            continue
        if all(ln.lstrip().startswith(("-", "*")) for ln in lines):
            for ln in lines:
                doc.add_bullet(ln.lstrip().lstrip("-*").strip())
        else:
            _add_paragraph(doc, " ".join(ln.strip() for ln in lines))


def re_split_blank(text):
    import re

    return [b for b in re.split(r"\n\s*\n", text) if b.strip()]


def _add_bullet(doc, text):
    # python-docx has no add_bullet; use list style.
    try:
        doc.add_paragraph(text, style="List Bullet")
    except KeyError:
        doc.add_paragraph("• " + text)


def build(research_path, out_path):
    data = _load(research_path)
    out_path = Path(out_path).resolve()
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Title block -----------------------------------------------------------
    title = doc.add_heading(data.get("title", "Blazepod — Deep Dive"), level=0)

    meta = doc.add_paragraph()
    meta.add_run("Generated: ").bold = True
    meta.add_run(data.get("generated_at") or datetime.now().isoformat())
    meta.add_run("\nSource: WAT workflow `blazepod-clone-research`").italic = True

    # Legend for confidence colors -----------------------------------------
    legend = doc.add_paragraph()
    legend.add_run("Confidence legend: ").bold = True
    for key in ("confirmed", "inferred", "unknown"):
        r = legend.add_run(f"{key} ")
        r.font.color.rgb = CONFIDENCE_COLORS[key]
        r.font.size = Pt(9)

    # Overview --------------------------------------------------------------
    _add_heading(doc, "1. Overview", level=1)
    _render_markdownish(doc, data.get("overview", ""))

    # Hardware --------------------------------------------------------------
    _add_heading(doc, "2. How the Hardware Works (explained simply)", level=1)
    _add_paragraph(
        doc,
        "This section explains each part of the physical pod in plain language, "
        "assuming no background in electronics. For each part we give the "
        "Blazepod spec, what it does, and a part you could use to build your own.",
        italic=True,
    )
    for i, hw in enumerate(data.get("hardware", []), start=1):
        _add_heading(doc, f"2.{i} {hw.get('subsystem', '(unnamed)')}", level=2)
        _add_confidence(doc, hw.get("confidence"))
        if hw.get("blazepod_spec"):
            _add_paragraph(doc, "Blazepod spec:", bold=True)
            _render_markdownish(doc, hw["blazepod_spec"])
        if hw.get("plain_explanation"):
            _add_paragraph(doc, "What it does (in plain terms):", bold=True)
            _render_markdownish(doc, hw["plain_explanation"])
        if hw.get("clone_part"):
            _add_paragraph(doc, "Clone part to use:", bold=True)
            _render_markdownish(doc, hw["clone_part"])

    # Software --------------------------------------------------------------
    _add_heading(doc, "3. How the Software Works", level=1)
    _add_paragraph(
        doc,
        "This section is the priority: how the app controls the pods and how "
        "the system behaves under software control.",
        italic=True,
    )
    for i, sw in enumerate(data.get("software", []), start=1):
        _add_heading(doc, f"3.{i} {sw.get('topic', '(topic)')}", level=2)
        _add_confidence(doc, sw.get("confidence"))
        _render_markdownish(doc, sw.get("explanation", ""))

    # Build notes -----------------------------------------------------------
    _add_heading(doc, "4. Build Notes — Mapping Subsystems to Clone Parts", level=1)
    build_notes = data.get("build_notes")
    if isinstance(build_notes, list):
        for note in build_notes:
            _add_heading(doc, note.get("subsystem", ""), level=2)
            _render_markdownish(doc, note.get("text", ""))
    elif isinstance(build_notes, str):
        _render_markdownish(doc, build_notes)
    else:
        _render_markdownish(
            doc,
            "See the companion BOM spreadsheet (blazepod-clone-bom.xlsx) for the "
            "full component mapping with Iranian sourcing.",
        )

    # Sources ---------------------------------------------------------------
    sources = data.get("sources", [])
    if sources:
        _add_heading(doc, "5. Sources", level=1)
        for s in sources:
            p = doc.add_paragraph(style="List Bullet")
            if s.get("url"):
                run = p.add_run(s.get("label", s["url"]))
                run.font.color.rgb = RGBColor(0x1A, 0x4E, 0x9E)
                run.font.underline = True
            else:
                p.add_run(s.get("label", ""))

    doc.save(out_path)
    sys.stdout.write(f"[build_doc] Wrote {out_path}\n")
    return out_path


def main():
    p = argparse.ArgumentParser(description=__doc__.splitlines()[1])
    p.add_argument("--research", required=True, help="Path to blazepod_research.json")
    p.add_argument("--out", required=True, help="Output .docx path")
    args = p.parse_args()
    build(args.research, args.out)


if __name__ == "__main__":
    main()
